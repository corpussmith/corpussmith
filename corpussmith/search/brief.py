"""Stage 10 — research-brief extraction from PDF, DOCX, MD, or plain text."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Tuple

SUPPORTED_EXTS = {".txt", ".md", ".markdown", ".rst", ".docx", ".pdf"}


class BriefExtractionError(Exception):
    pass


@dataclass
class Brief:
    title: str
    anchor_sections: List[Tuple[str, str]]  # (name_lower, body_text)
    seed_text: str


# ─── readers ──────────────────────────────────────────────────────────────────

def _read_raw(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in (".txt", ".md", ".markdown", ".rst"):
        return path.read_text(encoding="utf-8", errors="replace")
    if ext == ".docx":
        try:
            import docx  # type: ignore
            doc = docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            raise BriefExtractionError(
                "python-docx not installed; run: pip install python-docx"
            )
    if ext == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError:
            try:
                from PyPDF2 import PdfReader  # type: ignore  # noqa: N813
            except ImportError:
                raise BriefExtractionError(
                    "pypdf not installed; run: pip install pypdf"
                )
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    raise BriefExtractionError(f"unsupported extension: {path.suffix!r}")


def _is_caps_heading(line: str) -> bool:
    s = line.strip()
    if len(s) < 3:
        return False
    alpha = [c for c in s if c.isalpha()]
    return bool(alpha) and sum(c.isupper() for c in alpha) / len(alpha) >= 0.85


def _build_seed(title: str, anchors: List[Tuple[str, str]], fallback_text: str) -> str:
    cap = 600
    if not title:
        return fallback_text[:cap]
    parts = [title]
    remaining = cap - len(title)
    for _name, body in anchors:
        if remaining <= 5:
            break
        chunk = f"\n\n{body}"
        parts.append(chunk[:remaining])
        remaining -= len(chunk)
    if remaining > 10 and not anchors:
        lines = [l.strip() for l in fallback_text.splitlines()
                 if l.strip() and l.strip() != title]
        extra = " ".join(lines[:5])
        parts.append(f"\n\n{extra}"[:remaining])
    return "".join(parts)[:cap]


# ─── format-specific parsers ──────────────────────────────────────────────────

def _parse_markdown(text: str) -> Brief:
    title = ""
    anchors: List[Tuple[str, str]] = []
    cur_head: str | None = None
    cur_body: List[str] = []

    for line in text.splitlines():
        if line.startswith("# ") and not title:
            title = line[2:].strip()
        elif line.startswith("## "):
            if cur_head is not None:
                anchors.append((cur_head, " ".join(cur_body).strip()))
            cur_head = line[3:].strip().lower()
            cur_body = []
        elif cur_head is not None:
            s = line.strip()
            if s and not s.startswith("#"):
                cur_body.append(s)

    if cur_head is not None:
        anchors.append((cur_head, " ".join(cur_body).strip()))

    if not title:
        for line in text.splitlines():
            if line.strip() and not line.startswith("#"):
                title = line.strip()
                break

    return Brief(title=title, anchor_sections=anchors,
                 seed_text=_build_seed(title, anchors, text))


def _parse_plain(text: str) -> Brief:
    lines = text.splitlines()
    title = ""
    anchors: List[Tuple[str, str]] = []
    cur_head: str | None = None
    cur_body: List[str] = []

    for line in lines:
        s = line.strip()
        if not s:
            continue
        if not title:
            title = s
            continue
        if _is_caps_heading(s):
            if cur_head is not None:
                anchors.append((cur_head, " ".join(cur_body).strip()))
            cur_head = s.lower()
            cur_body = []
        elif cur_head is not None:
            cur_body.append(s)

    if cur_head is not None:
        anchors.append((cur_head, " ".join(cur_body).strip()))

    return Brief(title=title, anchor_sections=anchors,
                 seed_text=_build_seed(title, anchors, text))


def _parse_docx(path: Path) -> Brief:
    try:
        import docx  # type: ignore
        doc = docx.Document(str(path))
    except ImportError:
        raw = _read_raw(path)
        return _parse_plain(raw)

    title = ""
    anchors: List[Tuple[str, str]] = []
    cur_head: str | None = None
    cur_body: List[str] = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading 1" in style and not title:
            title = t
        elif "heading 2" in style:
            if cur_head is not None:
                anchors.append((cur_head, " ".join(cur_body).strip()))
            cur_head = t.lower()
            cur_body = []
        elif cur_head is not None:
            cur_body.append(t)
        elif not title:
            title = t

    if cur_head is not None:
        anchors.append((cur_head, " ".join(cur_body).strip()))

    full_text = "\n".join(p.text for p in doc.paragraphs)
    return Brief(title=title, anchor_sections=anchors,
                 seed_text=_build_seed(title, anchors, full_text))


# ─── public API ───────────────────────────────────────────────────────────────

def extract(path: Path) -> Brief:
    if not path.exists():
        raise BriefExtractionError(f"file not found: {path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise BriefExtractionError(f"unsupported extension: {path.suffix!r}")

    if ext == ".docx":
        return _parse_docx(path)

    raw = _read_raw(path)
    if not raw.strip():
        raise BriefExtractionError("file is empty")

    if ext in (".md", ".markdown"):
        return _parse_markdown(raw)
    return _parse_plain(raw)


__all__ = ["Brief", "BriefExtractionError", "extract", "SUPPORTED_EXTS"]
